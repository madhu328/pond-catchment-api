#!/usr/bin/env python3
"""
generate_report_docx.py
-----------------------
Generates a beautifully formatted, editable Word Document (.docx)
for Assignment 1 - Phase 2: Pond Catchment Analysis Backend Report.
"""

import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def build_docx():
    doc = Document()
    
    # Page Setup - Standard Margins (1 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Styles & Colors
    PRIMARY_COLOR = RGBColor(14, 116, 144)   # Deep Cyan / Teal
    DARK_TEXT = RGBColor(30, 41, 59)        # Slate 800
    CODE_COLOR = RGBColor(15, 23, 42)        # Slate 900
    ACCENT_COLOR = RGBColor(2, 132, 199)     # Sky Blue

    # Set normal font
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Segoe UI'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = DARK_TEXT

    # Document Header Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("Assignment 1 - Phase 2: Pond Catchment Analysis Backend Report")
    run_title.font.name = 'Segoe UI'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY_COLOR
    title.paragraph_format.space_after = Pt(12)

    # Metadata Card (Table)
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Author / Student:", "Madhu"),
        ("Instructor:", "Shivam Kushwaha"),
        ("Submission Date:", "September 2, 2026"),
        ("Course Assignment:", "Assignment 1 - Phase 2: Pond Catchment Analysis Backend"),
        ("GitHub Repository:", "https://github.com/madhu328/pond-catchment-api")
    ]
    
    for idx, (label, val) in enumerate(meta_data):
        row = meta_table.rows[idx]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        
        cell_lbl.width = Inches(2.0)
        cell_val.width = Inches(4.5)
        
        p_lbl = cell_lbl.paragraphs[0]
        r1 = p_lbl.add_run(label)
        r1.bold = True
        r1.font.size = Pt(10.5)
        r1.font.color.rgb = PRIMARY_COLOR
        
        p_val = cell_val.paragraphs[0]
        r2 = p_val.add_run(val)
        r2.font.size = Pt(10.5)
        
        set_cell_background(cell_lbl, "F1F5F9")
        set_cell_background(cell_val, "F8FAFC")
        set_cell_margins(cell_lbl, top=80, bottom=80, left=120, right=120)
        set_cell_margins(cell_val, top=80, bottom=80, left=120, right=120)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    def add_h1(text):
        h = doc.add_paragraph()
        r = h.add_run(text)
        r.font.name = 'Segoe UI'
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = PRIMARY_COLOR
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        return h

    def add_h2(text):
        h = doc.add_paragraph()
        r = h.add_run(text)
        r.font.name = 'Segoe UI'
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = ACCENT_COLOR
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)
        return h

    def add_code_block(code_text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.rows[0].cells[0]
        cell.width = Inches(6.5)
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(code_text.strip())
        r.font.name = 'Consolas'
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(74, 222, 128) # Bright Green
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 1. Project & API Route Overview
    add_h1("1. Project & API Route Overview")
    
    overview_table = doc.add_table(rows=9, cols=2)
    overview_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    overview_rows = [
        ("GitHub Repository", "https://github.com/madhu328/pond-catchment-api"),
        ("Primary API Endpoint (Network / TA)", "http://10.50.33.238:3313/analyzeContour"),
        ("Local Machine Endpoint", "http://localhost:3313/analyzeContour"),
        ("Public WAN Endpoint", "http://103.147.138.252:3313/analyzeContour"),
        ("Alias API Endpoint", "http://10.50.33.238:3313/findCatchment"),
        ("Required Form Field", "contour_map (File upload: .kml or .kmz)"),
        ("HTTP Method", "POST"),
        ("Interactive Docs (Swagger UI)", "http://10.50.33.238:3313/docs (or http://localhost:3313/docs)"),
        ("Interactive Upload Interface", "http://10.50.33.238:3313/ (or http://localhost:3313/)")
    ]

    for idx, (k, v) in enumerate(overview_rows):
        row = overview_table.rows[idx]
        cell_k, cell_v = row.cells[0], row.cells[1]
        cell_k.width = Inches(2.8)
        cell_v.width = Inches(3.7)
        
        pk = cell_k.paragraphs[0]
        rk = pk.add_run(k)
        rk.bold = True
        rk.font.size = Pt(10)
        
        pv = cell_v.paragraphs[0]
        rv = pv.add_run(v)
        rv.font.size = Pt(10)
        
        bg = "E2E8F0" if idx % 2 == 0 else "F8FAFC"
        set_cell_background(cell_k, bg)
        set_cell_background(cell_v, "FFFFFF")
        set_cell_margins(cell_k, top=60, bottom=60, left=100, right=100)
        set_cell_margins(cell_v, top=60, bottom=60, left=100, right=100)

    # 2. Catchment Estimation Approach & Methodology
    add_h1("2. Catchment Estimation Approach & Methodology")
    
    p_approach = doc.add_paragraph()
    p_approach.add_run(
        "The backend implements a fully generalized hydrological & terrain analysis pipeline "
        "that dynamically parses coordinate and elevation data from KML/KMZ files without relying "
        "on hardcoded coordinates or bounding boxes."
    )
    p_approach.paragraph_format.space_after = Pt(8)

    # ASCII Workflow Box
    workflow_ascii = """+-------------------------------------------------------+
|   Uploaded KML/KMZ File ('contour_map' form field)    |
+-------------------------------------------------------+
                           |
                           v
+-------------------------------------------------------+
|  1. Robust KML/KMZ Parser (XML / zipfile parsing)     |
|     - Extracts 3D coordinates (Lon, Lat, Elevation)   |
+-------------------------------------------------------+
                           |
                           v
+-------------------------------------------------------+
|  2. DEM Builder (SciPy Grid Interpolation)            |
|     - Projects points & generates elevation grid matrix|
+-------------------------------------------------------+
                           |
                           v
+-------------------------------------------------------+
|  3. D8 Flow Direction & Flow Accumulation Algorithm   |
|     - Calculates steepest slope & upstream drainage   |
+-------------------------------------------------------+
                           |
                           v
+-------------------------------------------------------+
|  4. Pond Site Selection & Catchment Delineation       |
|     - Selects natural convergence point               |
|     - Backtracks contributing upstream area           |
+-------------------------------------------------------+
                           |
                           v
+-------------------------------------------------------+
|  5. Structured JSON Output Generation                 |
|     - Lon/Lat, Elevation, Area (m² & Ha), Boundary    |
+-------------------------------------------------------+"""
    add_code_block(workflow_ascii)

    add_h2("Key Technical Steps")
    
    steps = [
        ("1. Dynamic KML/KMZ Parsing (app/kml_parser.py): ", 
         "Parses XML placemarks, LineStrings, and Polygons. Automatically unzips .kmz files to locate doc.kml. Extracts all 3D tuples (Longitude, Latitude, Elevation in meters)."),
        ("2. Digital Elevation Model (DEM) Interpolation (app/dem_builder.py): ", 
         "Converts geographic coordinates into an equidistant local meter-based grid. Uses SciPy griddata (Cubic & Nearest-Neighbor interpolation) to build a continuous regular raster grid of elevation values."),
        ("3. D8 Flow Routing & Accumulation (app/catchment.py): ", 
         "Computes the direction of steepest descent from each cell to its 8 neighbors (D8 Flow Direction). Resolves topological sorting to calculate Flow Accumulation—the exact number of upstream grid cells draining into each individual cell."),
        ("4. Optimal Pond Placement (app/catchment.py): ", 
         "Identifies cells combining high flow accumulation (high runoff collection potential) with low relative elevation. Selects the optimal cell for village pond construction."),
        ("5. Catchment Delineation & Area Calculation: ", 
         "Performs a recursive reverse traversal from the target pond cell to identify all contributing cells. Calculates exact catchment area: Area (m²) = Contributing Cells × (Grid Resolution)².")
    ]

    for title_text, desc_text in steps:
        p_step = doc.add_paragraph()
        r_t = p_step.add_run(title_text)
        r_t.bold = True
        r_t.font.color.rgb = ACCENT_COLOR
        p_step.add_run(desc_text)
        p_step.paragraph_format.space_after = Pt(6)

    # 3. Demonstration on Provided Sample Map (contours_1m.kml)
    add_h1("3. Demonstration on Provided Sample Map (contours_1m.kml)")

    add_h2("Execution Details")
    exec_p = doc.add_paragraph()
    exec_p.add_run("• Input File: ").bold = True
    exec_p.add_run("contours_1m.kml (6.71 MB)\n")
    exec_p.add_run("• Processing Time: ").bold = True
    exec_p.add_run("~2.16 seconds\n")
    exec_p.add_run("• Grid Resolution: ").bold = True
    exec_p.add_run("10.0 meters/cell\n")
    exec_p.add_run("• Grid Dimensions: ").bold = True
    exec_p.add_run("264 × 324 grid cells")
    exec_p.paragraph_format.space_after = Pt(8)

    add_h2("Returned Results Summary")
    res_p = doc.add_paragraph()
    res_p.add_run("Recommended Pond Site Location:\n").bold = True
    res_p.add_run("  - Latitude: 21.246056° N\n  - Longitude: 81.289336° E\n  - Site Elevation: 268.0 meters\n\n")
    res_p.add_run("Catchment Hydrology:\n").bold = True
    res_p.add_run("  - Total Catchment Area: 35,700.0 m² (3.57 Hectares)\n  - Contributing Cells: 357 cells\n  - Minimum Elevation: 268.0 m\n  - Maximum Elevation: 284.88 m")
    res_p.paragraph_format.space_after = Pt(8)

    # 4. API Specification & Documentation
    add_h1("4. API Specification & Documentation")

    add_h2("cURL Command Example")
    curl_str = """curl -X POST "http://10.50.33.238:3313/analyzeContour" \\
  -F "contour_map=@contours_1m.kml" """
    add_code_block(curl_str)

    add_h2("Postman Setup Instructions")
    postman_p = doc.add_paragraph()
    postman_p.add_run("1. Set HTTP Method to ").font.color.rgb = DARK_TEXT
    postman_p.add_run("POST").bold = True
    postman_p.add_run(".\n2. Enter URL: ")
    postman_p.add_run("http://10.50.33.238:3313/analyzeContour").bold = True
    postman_p.add_run(".\n3. Under the ")
    postman_p.add_run("Body").bold = True
    postman_p.add_run(" tab, select ")
    postman_p.add_run("form-data").bold = True
    postman_p.add_run(".\n4. Set Key = ")
    postman_p.add_run("contour_map").bold = True
    postman_p.add_run(" (change type dropdown from Text to File).\n5. Attach ")
    postman_p.add_run("contours_1m.kml").bold = True
    postman_p.add_run(".\n6. Click ")
    postman_p.add_run("Send").bold = True
    postman_p.paragraph_format.space_after = Pt(8)

    add_h2("Sample Output Response (JSON)")
    sample_json = """{
  "input_file": "contours_1m.kml",
  "processing_time_seconds": 2.16,
  "dem_resolution_meters": 10.0,
  "grid_size": {
    "rows": 264,
    "cols": 324
  },
  "recommended_pond_location": {
    "longitude": 81.28933599489764,
    "latitude": 21.24605558428102,
    "elevation_m": 268.0
  },
  "catchment": {
    "area_m2": 35700.0,
    "area_hectares": 3.57,
    "num_contributing_cells": 357,
    "elevation_min_m": 268.0,
    "elevation_max_m": 284.88,
    "boundary_polygon_lonlat": [
      [81.28846546444733, 21.24542323664919],
      [81.28807856202498, 21.245603907401144],
      [81.28788511081379, 21.24569424277712],
      [81.2877883852082, 21.245784578153092],
      [81.28759493399703, 21.245965248905044],
      [81.28749820839144, 21.246416925784924],
      [81.28769165960261, 21.246958938040777],
      [81.28807856202498, 21.247410614920657],
      [81.28827201323615, 21.247591285672605],
      [81.2887556412641, 21.247591285672605],
      [81.28894909247528, 21.247410614920657],
      [81.28933599489764, 21.247049273416753],
      [81.2896261717144, 21.24668793191285],
      [81.28933599489764, 21.24605558428102],
      [81.28856219005291, 21.24542323664919],
      [81.28846546444733, 21.24542323664919]
    ]
  },
  "note": "Runoff volume and rainfall-based storage estimates will be integrated in the next phase using a rainfall API, using this catchment area as the input."
}"""
    add_code_block(sample_json)

    # 5. Code Extensibility & Future Phase Readiness
    add_h1("5. Code Extensibility & Future Phase Readiness")
    ext_p = doc.add_paragraph()
    ext_p.add_run("1. Modular Architecture:\n").bold = True
    ext_p.add_run("   • app/kml_parser.py: Standalone parser supporting any standard KML/KMZ containing point, linestring, or polygon contour elements.\n")
    ext_p.add_run("   • app/dem_builder.py: Independent DEM generator scalable to custom grid resolutions.\n")
    ext_p.add_run("   • app/catchment.py: Pure mathematical D8 hydrology engine.\n")
    ext_p.add_run("   • app/main.py: Clean RESTful API wrapper using FastAPI.\n\n")
    ext_p.add_run("2. Phase 3 Integration Readiness:\n").bold = True
    ext_p.add_run("   • The calculated catchment area (35,700.0 m² / 3.57 ha) directly serves as the input parameter for precipitation runoff equations (Q = C · I · A).\n")
    ext_p.add_run("   • Future rainfall APIs (Open-Meteo / IMD) can easily query historical or real-time precipitation data (I) using the recommended_pond_location coordinates (21.246056° N, 81.289336° E).")
    ext_p.paragraph_format.space_after = Pt(12)

    # Save to file
    out_path = "/home/madhu/Desktop/pond_backend/Pond_Catchment_Analysis_Report.docx"
    doc.save(out_path)
    print(f"Successfully generated DOCX report at: {out_path}")

if __name__ == "__main__":
    build_docx()
